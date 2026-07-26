#!/usr/bin/env python3
from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MATERIALS_ROOT = ROOT / "@Materi Kuliah"
RPS_ROOT = Path(
    "/Users/mindra/Library/CloudStorage/GoogleDrive-mindra@unpad.ac.id/"
    "My Drive/@DOKUMEN S2/@AMI/RPS 2025"
)
BEGIN = "<!-- BEGIN ADOPSI AI RPS 2026 -->"
END = "<!-- END ADOPSI AI RPS 2026 -->"

RPS_BY_FOLDER = {
    "Materi_ADW_Tingkat_Lanjut": "RPS OBE 2025 Matakuliah Wajib/Analisis Deret Waktu Tingkat Lanjut/RPS Analisis Deret Waktu Tingkat Lanjut.docx",
    "Materi_Aktuaria_1": "RPS OBE 2025 Matakuliah Pilihan/Matematika Aktuaria 1/RPS Matematika Aktuaria 1.docx",
    "Materi_Analisis_Image": "RPS OBE 2025 Matakuliah Pilihan/Analisis Image/RPS Analisis Image.docx",
    "Materi_Analisis_Multilevel_dan_Longitudinal": "RPS OBE 2025 Matakuliah Pilihan/Analisis Multilevel dan Logitudinal/RPS Analisis Multilevel dan Longitudinal.docx",
    "Materi_Analisis_Multivariat_Tingkat_Lanjut": "RPS OBE 2025 Matakuliah Wajib/Analisis Multivariat Tingkat Lanjut/RPS Analisis Multivariat Tingkat Lanjut.docx",
    "Materi_Analisis_Regresi_Tingkat_Lanjut": "RPS OBE 2025 Matakuliah Wajib/Analisis Regresi Tingkat Lanjut/RPS Analisis Regresi Tingkat Lanjut.docx",
    "Materi_Analisis_Spasial": "RPS OBE 2025 Matakuliah Pilihan/Analisis Spasial/RPS Analisis Spasial.docx",
    "Materi_Analisis_Survival": "RPS OBE 2025 Matakuliah Pilihan/Analisis Survival/RPS Analisis Survival.docx",
    "Materi_Analisis_Teks": "RPS OBE 2025 Matakuliah Pilihan/Analisis Teks/RPS Analisis Teks.docx",
    "Materi_Basis_Data": "RPS OBE 2025 Matakuliah Pilihan/Basis data/RPS Basis data.docx",
    "Materi_Desain_Eksperimen": "RPS OBE 2025 Matakuliah Pilihan/Desain Eksperimen/RPS Desain Eksperimen.docx",
    "Materi_Epidemiologi": "RPS OBE 2025 Matakuliah Pilihan/Epidemiologi/RPS Epidemiologi.docx",
    "Materi_Komputasi_Statistik_dan_Optimasi": "RPS OBE 2025 Matakuliah Wajib/Komputasi Statistik dan Optimasi/RPS Komputasi Statistik dan Optimasi.docx",
    "Materi_Matematika_Aktuaria_2": "RPS OBE 2025 Matakuliah Pilihan/Matematika Aktuaria 2/RPS Matematika Aktuaria 2.docx",
    "Materi_Matematika_Keuangan": "RPS OBE 2025 Matakuliah Pilihan/Matematika Keuangan/RPS Matematika Keuangan.docx",
    "Materi_Model_Linear_Generalisasi": "RPS OBE 2025 Matakuliah Pilihan/Model Linear Generalisasi/RPS Model Linear Generalisasi.docx",
    "Materi_Pembelajaran_Mesin": "RPS OBE 2025 Matakuliah Pilihan/Pembelajaran Mesin/RPS Pembelajaran Mesin.docx",
    "Materi_Penambangan_Data_Kecerdasan_Buatan": "RPS OBE 2025 Matakuliah Wajib/Penambangan Data dan Kecerdasan Buatan/RPS Penambangan Data dan Kecerdasan Buatan.docx",
    "Materi_Proses_Stokastik_Tingkat_Lanjut": "RPS OBE 2025 Matakuliah Wajib/Proses Stokastik Tingkat Lanjut/RPS Proses Stokastik Tingkat Lanjut.docx",
    "Materi_Sampling_Survey": "RPS OBE 2025 Matakuliah Pilihan/Sampling Survey/RPS Sampling Survey.docx",
    "Materi_SEM": "RPS OBE 2025 Matakuliah Pilihan/Model Persamaan Struktural/RPS Pemodelan Persamaan Struktural.docx",
    "Materi_Statistika_Inferensial": "RPS OBE 2025 Matakuliah Wajib/Statistika Inferensial/RPS Statistik Inferensial.docx",
    "Materi_Statistika_Nonparametrik_Pemodelan_Fleksibel": "RPS OBE 2025 Matakuliah Wajib/Statistika Nonparametrik dan Pemodelan Fleksibel/RPS Statistika Nonparametrik dan Pemodelan Fleksibel.docx",
    "Materi_Teori_Antrian": "RPS OBE 2025 Matakuliah Pilihan/Teori Antrian/RPS Teori Antrian.docx",
    "Materi_Teori_Risiko": "RPS OBE 2025 Matakuliah Pilihan/Teori Risiko/RPS Teori Risiko.docx",
}


def unique(items):
    out = []
    for item in items:
        item = re.sub(r"\s+", " ", item).strip()
        if item and item not in out:
            out.append(item)
    return out


def extract_alignment(rps_path: Path):
    doc = Document(rps_path)
    course = rps_path.stem.removeprefix("RPS ").strip()
    cpmk = []
    ai_focus = []
    for table in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in table.rows]
        if rows and "Integrasi AI kontekstual" in rows[0]:
            header = rows[0]
            ai_idx = header.index("Integrasi AI kontekstual")
            for row in rows[1:]:
                if len(row) > ai_idx:
                    text = row[ai_idx]
                    if text and "Penggunaan AI hanya" not in text:
                        ai_focus.append(text)
        for row in rows:
            joined = " ".join(row)
            m = re.search(r"(CPMK\d+)\s*[:/|-]?\s*(Mampu .+)", joined)
            if m:
                statement = re.split(r"\.\s+", m.group(2), maxsplit=1)[0].rstrip(".") + "."
                cpmk.append(f"{m.group(1)}: {statement}")
    ai_focus = unique(ai_focus)[:3]
    cpmk = unique(cpmk)[:4]
    if not ai_focus:
        ai_focus = [
            f"AI digunakan untuk mempercepat eksplorasi, komputasi, dan pemeriksaan awal pada {course}.",
            "Hasil AI dibandingkan dengan derivasi, kode, diagnostik, dan sumber primer.",
            "Kesimpulan tetap ditetapkan mahasiswa melalui penalaran statistika dan konteks substantif.",
        ]
    return course, cpmk, ai_focus


def rmd_block(course, ai_focus):
    bullets = "\n".join(f"- {x}" for x in ai_focus)
    return f"""
{BEGIN}

## Penyelarasan Materi dengan RPS Revisi AI 2026

Materi **{course}** ini telah diselaraskan dengan RPS Revisi AI 2026. Rumusan CPL
program studi tetap dan tidak diubah. Generative AI maupun agentic AI hanya
digunakan sebagai alat bantu untuk mempercepat ketercapaian CPMK dan Sub-CPMK;
penalaran statistika, keputusan metodologis, dan tanggung jawab akademik tetap
berada pada mahasiswa.

### Integrasi AI kontekstual

{bullets}

### Protokol penggunaan AI

1. Nyatakan alat, tujuan penggunaan, dan bagian pekerjaan yang dibantu AI.
2. Simpan prompt log ringkas, sumber, versi kode/notebook, dan perubahan yang dilakukan.
3. Verifikasi keluaran terhadap asumsi metode, sumber primer, hasil numerik,
   diagnostik, ketidakpastian, robustness, bias/fairness, privasi, dan konteks.
4. Jangan memasukkan data pribadi, rahasia, soal asesmen tertutup, atau
   kredensial ke layanan AI yang tidak disetujui.
5. Mahasiswa wajib mampu menjelaskan dan mempertahankan seluruh hasil secara mandiri.

### Bukti asesmen

Penilaian berfokus pada ketepatan metode, validasi dan ketidakpastian,
reproducibility, interpretasi, komunikasi, serta transparansi penggunaan AI.
Keluaran AI tanpa verifikasi bukan bukti ketercapaian CPMK.

{END}
""".strip()


def html_block(course, ai_focus):
    lis = "".join(f"<li>{x}</li>" for x in ai_focus)
    return f"""
{BEGIN}
<section id="penyelarasan-rps-ai-2026" style="margin:2rem 0;padding:1.5rem;border-left:8px solid #f2a51a;background:#eef5f8;border-radius:8px">
  <h2>Penyelarasan Materi dengan RPS Revisi AI 2026</h2>
  <p>Materi <strong>{course}</strong> ini mempertahankan rumusan CPL yang berlaku. AI ditempatkan sebagai alat bantu untuk mempercepat ketercapaian CPMK dan Sub-CPMK, bukan sebagai CPL baru atau pengganti penalaran statistika.</p>
  <h3>Integrasi AI kontekstual</h3><ul>{lis}</ul>
  <h3>Protokol penggunaan AI</h3>
  <ol>
    <li>Ungkapkan alat, tujuan, prompt log ringkas, sumber, dan perubahan yang dilakukan.</li>
    <li>Verifikasi asumsi, kode, hasil numerik, diagnostik, ketidakpastian, robustness, bias/fairness, privasi, dan konteks.</li>
    <li>Jangan mengunggah data pribadi, rahasia, asesmen tertutup, atau kredensial ke layanan yang tidak disetujui.</li>
    <li>Mahasiswa wajib mampu menjelaskan dan mempertahankan seluruh hasil secara mandiri.</li>
  </ol>
  <p><strong>Bukti asesmen:</strong> ketepatan metode, validasi, reproducibility, interpretasi, komunikasi, dan transparansi penggunaan AI. Keluaran AI tanpa verifikasi bukan bukti ketercapaian CPMK.</p>
</section>
{END}
""".strip()


def replace_marked(text, block):
    pattern = re.compile(re.escape(BEGIN) + r".*?" + re.escape(END), re.S)
    if pattern.search(text):
        return pattern.sub(block, text)
    return text.rstrip() + "\n\n" + block + "\n"


def revise_rmd(path, course, ai_focus):
    text = path.read_text(encoding="utf-8")
    path.write_text(replace_marked(text, rmd_block(course, ai_focus)), encoding="utf-8")


def revise_html(path, course, ai_focus):
    text = path.read_text(encoding="utf-8")
    text = re.sub(re.escape(BEGIN) + r".*?" + re.escape(END), "", text, flags=re.S)
    block = html_block(course, ai_focus)
    marker = "</div>\n\n<script>"
    if marker in text:
        text = text.replace(marker, block + "\n</div>\n\n<script>", 1)
    elif "</body>" in text:
        text = text.replace("</body>", block + "\n</body>", 1)
    else:
        text += "\n" + block
    path.write_text(text, encoding="utf-8")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def revise_contract(path, course, cpmk, ai_focus):
    doc = Document(path)
    for p in doc.paragraphs:
        for run in p.runs:
            run.text = run.text.replace("Januari 2025", "Juli 2026").replace("Tahun 2025", "Tahun 2026")
    # Make the operation idempotent by removing a prior generated tail.
    marker = "PENYELARASAN KONTRAK DENGAN RPS REVISI AI 2026"
    for p in list(doc.paragraphs):
        if p.text.strip() == marker:
            body = p._element.getparent()
            idx = body.index(p._element)
            for element in list(body)[idx:]:
                body.remove(element)
            break
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 1"] if "Heading 1" in doc.styles else doc.styles["Normal"]
    run = heading.add_run(marker)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(23, 57, 92)
    doc.add_paragraph(
        f"Kontrak perkuliahan {course} ini mengikuti RPS Revisi AI 2026. "
        "Rumusan CPL tidak diubah. AI berfungsi sebagai alat bantu untuk "
        "mempercepat ketercapaian CPMK/Sub-CPMK, bukan pengganti penalaran "
        "statistika, keputusan metodologis, atau tanggung jawab akademik."
    )
    if cpmk:
        doc.add_heading("CPMK yang dipercepat melalui dukungan AI", level=2)
        for item in cpmk:
            doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Integrasi AI kontekstual", level=2)
    for item in ai_focus:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Aturan penggunaan dan bukti yang wajib diserahkan", level=2)
    rules = [
        "AI boleh digunakan untuk eksplorasi, penyusunan alternatif, pemeriksaan kode, simulasi, atau umpan balik awal sesuai instruksi dosen.",
        "Setiap penggunaan AI wajib disertai pernyataan penggunaan, prompt log ringkas, sumber, versi alat, perubahan yang dilakukan, dan bukti verifikasi mandiri.",
        "Keluaran wajib diperiksa terhadap asumsi metode, sumber primer, hasil numerik, diagnostik, ketidakpastian, robustness, bias/fairness, privasi, dan konteks.",
        "Data pribadi, rahasia, soal asesmen tertutup, materi berhak cipta tanpa izin, dan kredensial tidak boleh dimasukkan ke layanan AI yang tidak disetujui.",
        "Pada asesmen yang dinyatakan tanpa AI, mahasiswa wajib menaati pembatasan tersebut. Mahasiswa selalu wajib mampu menjelaskan seluruh hasil secara mandiri.",
    ]
    for rule in rules:
        doc.add_paragraph(rule, style="List Number")
    doc.add_heading("Rubrik minimum asesmen berbantuan AI", level=2)
    rubrics = [
        ("Ketepatan metodologi", "Asumsi, metode, derivasi/kode, dan inferensi benar serta sesuai konteks.", "35%"),
        ("Validasi dan ketidakpastian", "Hasil diuji ulang; diagnostik, robustness, ketidakpastian, dan keterbatasan dibahas.", "25%"),
        ("Reproducibility dan audit trail", "Data dictionary, kode/notebook, versi alat, prompt log, sumber, dan jejak agen memadai.", "20%"),
        ("Responsible AI dan komunikasi", "Privasi, fairness, transparansi, human oversight, dampak, dan komunikasi dipenuhi.", "20%"),
    ]
    for dimension, indicator, weight in rubrics:
        paragraph = doc.add_paragraph(style="List Bullet")
        label = paragraph.add_run(f"{dimension} ({weight}) — ")
        label.bold = True
        paragraph.add_run(indicator)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
    # python-docx may append after the terminal body-level sectPr in documents
    # with several sections. Word/LibreOffice ignores content in that position,
    # so move every generated trailing element immediately before sectPr.
    body = doc._element.body
    terminal_sect_pr = body.sectPr
    if terminal_sect_pr is not None:
        terminal_index = body.index(terminal_sect_pr)
        trailing = list(body)[terminal_index + 1 :]
        for element in trailing:
            body.remove(element)
            body.insert(terminal_index, element)
            terminal_index += 1
    doc.core_properties.modified = datetime.now()
    doc.save(path)


def main():
    catalog_path = ROOT / "data/materials.json"
    catalog = json.loads(catalog_path.read_text(encoding="utf-8"))
    report = []
    for material in catalog["materials"]:
        folder = material["folder"]
        directory = MATERIALS_ROOT / folder
        rps_path = RPS_ROOT / RPS_BY_FOLDER[folder]
        if not rps_path.exists():
            raise FileNotFoundError(rps_path)
        course, cpmk, ai_focus = extract_alignment(rps_path)
        rmds = list(directory.glob("*.Rmd"))
        html_path = directory / material["file"]
        contract_path = directory / material["contractFile"]
        for rmd in rmds:
            revise_rmd(rmd, course, ai_focus)
        revise_html(html_path, course, ai_focus)
        revise_contract(contract_path, course, cpmk, ai_focus)
        material["revision"] = "RPS Revisi AI 2026"
        material["aiAdoption"] = True
        material["aiRole"] = "Akselerator ketercapaian CPMK/Sub-CPMK; CPL tetap"
        material["updatedAt"] = "2026-07-27"
        report.append(
            {
                "folder": folder,
                "course": course,
                "rps": str(rps_path),
                "rmd": [p.name for p in rmds],
                "html": html_path.name,
                "contract": contract_path.name,
                "aiFocus": ai_focus,
            }
        )
    catalog["generatedAt"] = "2026-07-27T00:00:00+07:00"
    catalog["revision"] = "RPS Revisi AI 2026"
    catalog["policy"] = "CPL tetap; AI mempercepat ketercapaian CPMK/Sub-CPMK"
    catalog_path.write_text(json.dumps(catalog, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    (ROOT / "data/materials-ai-revision-report.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    print(f"Revised {len(report)} material packages.")


if __name__ == "__main__":
    main()

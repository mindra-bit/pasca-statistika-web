#!/usr/bin/env python3
"""Restore course-contract covers and fix dark table header text."""

from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from pathlib import Path
import subprocess

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[1]
MATERIALS = ROOT / "@Materi Kuliah"
SOURCE_COMMIT = "2c51107"
DARK_HEADER_FILL = "102A43"


def relative_posix(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def historical_document(path: Path) -> Document:
    spec = f"{SOURCE_COMMIT}:{relative_posix(path)}"
    blob = subprocess.check_output(["git", "show", spec], cwd=ROOT)
    return Document(BytesIO(blob))


def restore_cover(current: Document, original: Document) -> None:
    if not original.inline_shapes:
        raise RuntimeError("Historical document has no cover image")

    old_para = original.paragraphs[0]
    blips = old_para._p.xpath(".//a:blip")
    if len(blips) != 1:
        raise RuntimeError(f"Expected one cover image, found {len(blips)}")

    old_rid = blips[0].get(qn("r:embed"))
    image_blob = original.part.related_parts[old_rid].blob
    shape = original.inline_shapes[0]

    current_para = current.paragraphs[0]
    for child in list(current_para._p):
        if child.tag != qn("w:pPr"):
            current_para._p.remove(child)

    current_para.alignment = old_para.alignment
    if old_para._p.pPr is not None:
        old_ppr = deepcopy(old_para._p.pPr)
        if current_para._p.pPr is not None:
            current_para._p.remove(current_para._p.pPr)
        current_para._p.insert(0, old_ppr)

    current_para.add_run().add_picture(
        BytesIO(image_blob), width=shape.width, height=shape.height
    )


def make_dark_headers_white(document: Document) -> tuple[int, int]:
    cells_changed = 0
    runs_changed = 0

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                shading = cell._tc.tcPr.find(qn("w:shd"))
                fill = (
                    shading.get(qn("w:fill"), "").upper()
                    if shading is not None
                    else ""
                )
                if fill != DARK_HEADER_FILL:
                    continue

                cells_changed += 1
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        runs_changed += 1

    return cells_changed, runs_changed


def repair(path: Path) -> tuple[int, int]:
    current = Document(path)
    original = historical_document(path)
    restore_cover(current, original)
    result = make_dark_headers_white(current)
    current.save(path)
    return result


def main() -> None:
    paths = sorted(MATERIALS.rglob("*Kontrak*.docx"))
    if len(paths) != 25:
        raise RuntimeError(f"Expected 25 course contracts, found {len(paths)}")

    total_cells = 0
    total_runs = 0
    for path in paths:
        cells, runs = repair(path)
        total_cells += cells
        total_runs += runs
        print(f"{relative_posix(path)}: cover restored; {cells} header cells fixed")

    print(
        f"Completed {len(paths)} contracts; "
        f"{total_cells} dark header cells and {total_runs} text runs updated."
    )


if __name__ == "__main__":
    main()
